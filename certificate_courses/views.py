# certificate_courses/views.py
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.urls import reverse
from django.http import HttpResponseForbidden, HttpResponse
from django.utils import timezone
from datetime import datetime
import os
from django.template.loader import get_template
from django.conf import settings
from io import BytesIO
import tempfile
from django.db import models

from .models import (
    CertificateCourse, Student, CourseEnrollment,
    AttendanceRecord, Certificate, CourseReport, CertificateTemplate
)
from .forms import (
    CertificateCourseForm, CourseOutcomeForm, StudentForm,
    CourseEnrollmentForm, AttendanceForm, StudentBulkImportForm, CertificateTemplateForm
)

from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage

# Certificate Course CRUD Operations
@login_required
def course_list(request):
    department = request.user.department
    courses = CertificateCourse.objects.filter(department=department)
    
    # Search functionality
    search_query = request.GET.get('search', '')
    if search_query:
        courses = courses.filter(
            models.Q(name__icontains=search_query) | 
            models.Q(course_code__icontains=search_query)
        )
    
    # Status filter
    status_filter = request.GET.get('status', '')
    if status_filter and status_filter in dict(CertificateCourse.STATUS_CHOICES).keys():
        courses = courses.filter(status=status_filter)
    
    # Date filter
    start_date = request.GET.get('start_date', '')
    end_date = request.GET.get('end_date', '')
    
    if start_date:
        try:
            start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
            courses = courses.filter(start_date__gte=start_date)
        except ValueError:
            messages.warning(request, 'Invalid start date format.')
    
    if end_date:
        try:
            end_date = datetime.strptime(end_date, '%Y-%m-%d').date()
            courses = courses.filter(end_date__lte=end_date)
        except ValueError:
            messages.warning(request, 'Invalid end date format.')
    
    # Pagination
    paginator = Paginator(courses, 10)  # Show 10 courses per page
    page_number = request.GET.get('page', 1)
    
    try:
        page_obj = paginator.page(page_number)
    except (PageNotAnInteger, EmptyPage):
        page_obj = paginator.page(1)
    
    context = {
        'courses': page_obj,
        'search_query': search_query,
        'status_filter': status_filter,
        'start_date': start_date if isinstance(start_date, str) else start_date.strftime('%Y-%m-%d'),
        'end_date': end_date if isinstance(end_date, str) else end_date.strftime('%Y-%m-%d'),
        'status_choices': CertificateCourse.STATUS_CHOICES,
    }
    
    return render(request, 'certificate_courses/course_list.html', context)

@login_required
def course_create(request):
    department = request.user.department
    
    if request.method == 'POST':
        form = CertificateCourseForm(request.POST, request.FILES)
        if form.is_valid():
            course = form.save(commit=False)
            course.department = department
            course.save()
            messages.success(request, 'Certificate course created successfully.')
            return redirect('course_detail', pk=course.pk)
    else:
        form = CertificateCourseForm()
    
    return render(request, 'certificate_courses/course_form.html', {'form': form})

@login_required
def course_detail(request, pk):
    department = request.user.department
    course = get_object_or_404(CertificateCourse, pk=pk)
    
    # Ensure the course belongs to the department
    if course.department != department:
        return HttpResponseForbidden("You don't have permission to access this course.")
    
    enrollments = course.enrollments.all()
    
    context = {
        'course': course,
        'enrollments': enrollments,
        'can_generate_report': course.can_generate_report(),
    }
    
    return render(request, 'certificate_courses/course_detail.html', context)

@login_required
def course_update(request, pk):
    department = request.user.department
    course = get_object_or_404(CertificateCourse, pk=pk)
    
    # Ensure the course belongs to the department
    if course.department != department:
        return HttpResponseForbidden("You don't have permission to access this course.")
    
    if request.method == 'POST':
        form = CertificateCourseForm(request.POST, request.FILES, instance=course)
        if form.is_valid():
            form.save()
            messages.success(request, 'Certificate course updated successfully.')
            return redirect('course_detail', pk=course.pk)
    else:
        form = CertificateCourseForm(instance=course)
    
    return render(request, 'certificate_courses/course_form.html', {'form': form, 'course': course})

@login_required
def course_delete(request, pk):
    department = request.user.department
    course = get_object_or_404(CertificateCourse, pk=pk)
    
    # Ensure the course belongs to the department
    if course.department != department:
        return HttpResponseForbidden("You don't have permission to access this course.")
    
    if request.method == 'POST':
        course.delete()
        messages.success(request, 'Certificate course deleted successfully.')
        return redirect('course_list')
    
    return render(request, 'certificate_courses/course_confirm_delete.html', {'course': course})

@login_required
def update_course_outcome(request, pk):
    department = request.user.department
    course = get_object_or_404(CertificateCourse, pk=pk)
    
    # Ensure the course belongs to the department
    if course.department != department:
        return HttpResponseForbidden("You don't have permission to access this course.")
    
    if request.method == 'POST':
        form = CourseOutcomeForm(request.POST, instance=course)
        if form.is_valid():
            form.save()
            messages.success(request, 'Course outcome updated successfully.')
            return redirect('course_detail', pk=course.pk)
    else:
        form = CourseOutcomeForm(instance=course)
    
    return render(request, 'certificate_courses/course_outcome_form.html', {'form': form, 'course': course})


# certificate_courses/views.py (continued)

# Student Management
@login_required
def student_list(request):
    department = request.user.department
    students = Student.objects.filter(department=department)
    return render(request, 'certificate_courses/student_list.html', {'students': students})

@login_required
def student_create(request):
    department = request.user.department
    
    if request.method == 'POST':
        form = StudentForm(request.POST)
        if form.is_valid():
            student = form.save(commit=False)
            student.department = department
            student.save()
            messages.success(request, 'Student created successfully.')
            return redirect('student_list')
    else:
        form = StudentForm()
    
    return render(request, 'certificate_courses/student_form.html', {'form': form})

@login_required
def student_update(request, pk):
    department = request.user.department
    student = get_object_or_404(Student, pk=pk, department=department)
    
    if request.method == 'POST':
        form = StudentForm(request.POST, instance=student)
        if form.is_valid():
            form.save()
            messages.success(request, 'Student updated successfully.')
            return redirect('student_list')
    else:
        form = StudentForm(instance=student)
    
    return render(request, 'certificate_courses/student_form.html', {'form': form, 'student': student})

@login_required
def student_delete(request, pk):
    department = request.user.department
    student = get_object_or_404(Student, pk=pk, department=department)
    
    if request.method == 'POST':
        student.delete()
        messages.success(request, 'Student deleted successfully.')
        return redirect('student_list')
    
    return render(request, 'certificate_courses/student_confirm_delete.html', {'student': student})

# Course Enrollment
@login_required
def manage_enrollment(request, course_pk):
    department = request.user.department
    course = get_object_or_404(CertificateCourse, pk=course_pk, department=department)
    
    if request.method == 'POST':
        form = CourseEnrollmentForm(request.POST, department=department)
        if form.is_valid():
            selected_students = form.cleaned_data['students']
            
            # Get existing enrollments
            existing_enrollments = CourseEnrollment.objects.filter(course=course)
            existing_student_ids = [e.student.id for e in existing_enrollments]
            
            # Add new enrollments
            for student in selected_students:
                if student.id not in existing_student_ids:
                    CourseEnrollment.objects.create(student=student, course=course)
            
            # Remove enrollments for unselected students
            for enrollment in existing_enrollments:
                if enrollment.student not in selected_students:
                    enrollment.delete()
            
            messages.success(request, 'Enrollments updated successfully.')
            return redirect('course_detail', pk=course.pk)
    else:
        # Pre-select currently enrolled students
        initial_students = Student.objects.filter(enrollments__course=course)
        form = CourseEnrollmentForm(department=department, initial={'students': initial_students})
    
    # The following return statement was missing
    return render(request, 'certificate_courses/manage_enrollment.html', {
        'form': form,
        'course': course,
        'enrolled_students': initial_students if 'initial_students' in locals() else []
    })


# certificate_courses/views.py (continued)

@login_required
def attendance_list(request, course_pk):
    department = request.user.department
    course = get_object_or_404(CertificateCourse, pk=course_pk, department=department)
    enrollments = course.enrollments.all()
    
    # Get all dates for which attendance was recorded
    attendance_dates = AttendanceRecord.objects.filter(
        enrollment__course=course
    ).values_list('date', flat=True).distinct().order_by('date')
    
    # Convert QuerySet to list for easier manipulation
    attendance_dates = list(attendance_dates)
    
    # Create attendance data structure
    attendance_data = []
    
    for enrollment in enrollments:
        present_count = 0
        attendance_records = []
        
        # Get attendance records for this enrollment
        records = AttendanceRecord.objects.filter(enrollment=enrollment)
        
        # Create a record for each date
        for date in attendance_dates:
            try:
                record = records.get(date=date)
                is_present = record.is_present
                if is_present:
                    present_count += 1
            except AttendanceRecord.DoesNotExist:
                is_present = False
                
            attendance_records.append({
                'date': date,
                'is_present': is_present
            })
        
        # Calculate attendance percentage
        attendance_percentage = 0
        if attendance_dates:
            attendance_percentage = int((present_count / len(attendance_dates)) * 100)
        
        student_data = {
            'student': enrollment.student,
            'enrollment': enrollment,
            'attendance_records': attendance_records,
            'attendance_percentage': attendance_percentage
        }
        
        attendance_data.append(student_data)
    
    context = {
        'course': course,
        'attendance_data': attendance_data,
        'attendance_dates': attendance_dates,
    }
    
    return render(request, 'certificate_courses/attendance_list.html', context)


@login_required
def take_attendance(request, course_pk):
    department = request.user.department
    course = get_object_or_404(CertificateCourse, pk=course_pk, department=department)
    enrollments = course.enrollments.all()
    
    if not enrollments:
        messages.warning(request, 'No students enrolled in this course yet.')
        return redirect('course_detail', pk=course_pk)
    
    if request.method == 'POST':
        date_str = request.POST.get('attendance_date')
        try:
            attendance_date = datetime.strptime(date_str, '%Y-%m-%d').date()
        except ValueError:
            messages.error(request, 'Invalid date format.')
            return redirect('take_attendance', course_pk=course_pk)
        
        form = AttendanceForm(request.POST, enrollments=enrollments, date=attendance_date)
        if form.is_valid():
            for enrollment in enrollments:
                field_name = f'attendance_{enrollment.id}'
                is_present = form.cleaned_data.get(field_name, False)
                
                # Update or create attendance record
                AttendanceRecord.objects.update_or_create(
                    enrollment=enrollment,
                    date=attendance_date,
                    defaults={'is_present': is_present}
                )
            
            messages.success(request, f'Attendance for {attendance_date} recorded successfully.')
            return redirect('attendance_list', course_pk=course_pk)
    else:
        # Default to today's date
        today = timezone.now().date()
        form = AttendanceForm(enrollments=enrollments, date=today)
    
    context = {
        'course': course,
        'form': form,
        'enrollments': enrollments,
        'attendance_date': request.GET.get('date', timezone.now().date().strftime('%Y-%m-%d')),
    }
    
    return render(request, 'certificate_courses/take_attendance.html', context)

# certificate_courses/views.py (continued)

# Certificate Generation
@login_required
def generate_certificates(request, course_pk):
    department = request.user.department
    course = get_object_or_404(CertificateCourse, pk=course_pk, department=department)
    
    # Check if course is completed
    if course.status != 'completed':
        messages.error(request, 'Cannot generate certificates for courses that are not completed.')
        return redirect('course_detail', pk=course_pk)
    
    # Check if course outcome is added
    if not course.course_outcome:
        messages.error(request, 'Please add course outcome before generating certificates.')
        return redirect('update_course_outcome', pk=course_pk)
    
    enrollments = course.enrollments.all()
    
    # Prepare enrollment data with attendance statistics
    enrollment_data = []
    
    for enrollment in enrollments:
        total_days = AttendanceRecord.objects.filter(enrollment=enrollment).count()
        present_days = AttendanceRecord.objects.filter(enrollment=enrollment, is_present=True).count()
        
        attendance_percentage = (present_days / total_days) * 100 if total_days > 0 else 0
        
        enrollment_data.append({
            'enrollment': enrollment,
            'total_days': total_days,
            'present_days': present_days,
            'attendance_percentage': attendance_percentage
        })
    
    if request.method == 'POST':
        # Process the selected enrollments
        selected_enrollments = request.POST.getlist('enrollments')
        
        for enrollment_id in selected_enrollments:
            enrollment = get_object_or_404(CourseEnrollment, pk=enrollment_id, course=course)
            
            # Check if student has sufficient attendance (example: 75% attendance required)
            total_days = AttendanceRecord.objects.filter(enrollment=enrollment).count()
            if total_days == 0:
                messages.warning(request, f'No attendance records for {enrollment.student.name}.')
                continue
                
            present_days = AttendanceRecord.objects.filter(enrollment=enrollment, is_present=True).count()
            attendance_percentage = (present_days / total_days) * 100
            
            if attendance_percentage < 75:
                messages.warning(
                    request, 
                    f'{enrollment.student.name} has attendance below 75% ({attendance_percentage:.1f}%).'
                )
                continue
            
            # Generate certificate
            # In a real application, you would use a library like ReportLab to generate PDF certificates
            # Here we'll create a dummy file for demonstration
            
            # Check if certificate already exists
            if hasattr(enrollment, 'certificate'):
                # Update existing certificate
                certificate = enrollment.certificate
            else:
                # Create new certificate
                certificate = Certificate(enrollment=enrollment)
            
            # For now, we'll just create a text file as placeholder
            # In a real app, you'd generate a proper PDF certificate
            certificate_content = f"""
            CERTIFICATE OF COMPLETION
            
            This is to certify that
            
            {enrollment.student.name}
            
            has successfully completed the course
            
            {course.name} ({course.course_code})
            
            Duration: {course.start_date} to {course.end_date}
            Total Hours: {course.total_hours}
            
            Issued on: {timezone.now().date()}
            """
            
            # Create a temporary file
            with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as temp_file:
                temp_file.write(certificate_content.encode('utf-8'))
                temp_path = temp_file.name
            
            # Save the certificate file to the model
            with open(temp_path, 'rb') as f:
                certificate_filename = f"certificate_{course.course_code}_{enrollment.student.registration_number}.txt"
                certificate.certificate_file.save(certificate_filename, BytesIO(f.read()))
            
            # Clean up the temporary file
            os.unlink(temp_path)
            
            messages.success(request, f'Certificate for {enrollment.student.name} generated successfully.')
        
        return redirect('course_detail', pk=course_pk)
    
    context = {
        'course': course,
        'enrollments': enrollments,
        'enrollment_data': enrollment_data,
    }
    
    return render(request, 'certificate_courses/generate_certificates.html', context)

@login_required
def view_certificate(request, certificate_pk):
    department = request.user.department
    certificate = get_object_or_404(Certificate, pk=certificate_pk, enrollment__course__department=department)
    
    # In a real application, you would serve the certificate file
    # Here we'll just redirect to the certificate file URL
    return redirect(certificate.certificate_file.url)


# certificate_courses/views.py (continued)

# Report Generation
# Report Generation
@login_required
def generate_report(request, course_pk):
    department = request.user.department
    course = get_object_or_404(CertificateCourse, pk=course_pk, department=department)
    enrollments = course.enrollments.all()
    can_generate_report = course.can_generate_report()
    
    # Check if report can be generated
    if not course.can_generate_report():
        messages.error(request, 'Cannot generate report. Please check if all requirements are met.')
        return redirect('course_detail', pk=course_pk)
    
    if request.method == 'POST':
        # Get the requested format (default to PDF if not specified)
        report_format = request.POST.get('report_format', 'pdf').lower()
        
        if report_format == 'docx':
            # Generate Word document using python-docx
            try:
                from docx import Document
                from docx.shared import Inches, Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                document = Document()
                
                # Add title
                title = document.add_heading('CERTIFICATE COURSE REPORT', level=0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add course information
                document.add_heading('Course Information', level=1)
                table = document.add_table(rows=5, cols=2)
                table.style = 'Table Grid'
                
                # Fill course info table
                rows = table.rows
                rows[0].cells[0].text = 'Course Name'
                rows[0].cells[1].text = course.name
                rows[1].cells[0].text = 'Course Code'
                rows[1].cells[1].text = course.course_code
                rows[2].cells[0].text = 'Department'
                rows[2].cells[1].text = course.department.name
                rows[3].cells[0].text = 'Duration'
                rows[3].cells[1].text = f"{course.start_date} to {course.end_date}"
                rows[4].cells[0].text = 'Total Hours'
                rows[4].cells[1].text = str(course.total_hours)
                
                # Course objective
                document.add_heading('Course Objective', level=1)
                document.add_paragraph(course.course_objective)
                
                # Course outcome
                document.add_heading('Course Outcome', level=1)
                document.add_paragraph(course.course_outcome)
                
                # Enrollment details
                document.add_heading('Enrollment Details', level=1)
                document.add_paragraph(f"Total Students Enrolled: {enrollments.count()}")
                
                # Student list with attendance
                document.add_heading('List of Students', level=1)
                student_table = document.add_table(rows=1, cols=4)
                student_table.style = 'Table Grid'
                
                # Add header row
                header_cells = student_table.rows[0].cells
                header_cells[0].text = 'Student Name'
                header_cells[1].text = 'Registration Number'
                header_cells[2].text = 'Attendance'
                header_cells[3].text = 'Certificate Issued'
                
                # Add student data
                for enrollment in enrollments:
                    row_cells = student_table.add_row().cells
                    
                    # Calculate attendance
                    total_days = AttendanceRecord.objects.filter(enrollment=enrollment).count()
                    present_days = AttendanceRecord.objects.filter(enrollment=enrollment, is_present=True).count()
                    attendance_percentage = (present_days / total_days) * 100 if total_days > 0 else 0
                    
                    row_cells[0].text = enrollment.student.name
                    row_cells[1].text = enrollment.student.registration_number
                    row_cells[2].text = f"{present_days}/{total_days} ({attendance_percentage:.1f}%)"
                    row_cells[3].text = 'Yes' if hasattr(enrollment, 'certificate') else 'No'
                
                # Add generation date
                document.add_paragraph(f"Report Generated On: {timezone.now().date()}")
                
                # Add signature lines
                document.add_paragraph()
                document.add_paragraph()
                signature_table = document.add_table(rows=1, cols=3)
                signature_cells = signature_table.rows[0].cells
                
                for i, title in enumerate(['Course Coordinator', 'HOD', 'Principal']):
                    signature_cells[i].text = '_'*30 + f"\n{title}"
                
                # Save to a temporary file
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                    temp_path = temp_file.name
                
                document.save(temp_path)
                
                # Save the report file to the model
                try:
                    report = course.report
                except CourseReport.DoesNotExist:
                    report = CourseReport(course=course)
                
                with open(temp_path, 'rb') as f:
                    report_filename = f"report_{course.course_code}.docx"
                    report.report_file.save(report_filename, BytesIO(f.read()))
                
                # Clean up the temporary file
                os.unlink(temp_path)
                
                messages.success(request, 'Report generated successfully as Word document.')
                
                # Return the file as a download
                response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename="{report_filename}"'
                with open(report.report_file.path, 'rb') as f:
                    response.write(f.read())
                return response
                
            except ImportError:
                messages.warning(request, 'Python-docx library not installed. Generating PDF instead.')
                report_format = 'pdf'
        
        # Generate PDF (default or fallback)
        if report_format == 'pdf':
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.lib import colors
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                
                # Create a temporary file to save the PDF to
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                    temp_path = temp_file.name
                
                # Create the PDF document
                doc = SimpleDocTemplate(
                    temp_path,
                    pagesize=letter,
                    rightMargin=72, leftMargin=72,
                    topMargin=72, bottomMargin=72
                )
                
                # Container for the elements to be added to the PDF
                elements = []
                
                # Styles
                styles = getSampleStyleSheet()
                styles.add(ParagraphStyle(name='Center', alignment=1))
                
                # Title
                elements.append(Paragraph('CERTIFICATE COURSE REPORT', styles['Title']))
                elements.append(Spacer(1, 12))
                
                # Course Information Table
                course_data = [
                    ['Course Name', course.name],
                    ['Course Code', course.course_code],
                    ['Department', course.department.name],
                    ['Duration', f"{course.start_date} to {course.end_date}"],
                    ['Total Hours', str(course.total_hours)]
                ]
                
                course_table = Table(course_data, colWidths=[120, 350])
                course_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                    ('TEXTCOLOR', (0, 0), (0, -1), colors.black),
                    ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (0, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                    ('TOPPADDING', (0, 0), (-1, -1), 12),
                    ('BACKGROUND', (1, 0), (-1, -1), colors.white),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                elements.append(Paragraph('Course Information', styles['Heading1']))
                elements.append(course_table)
                elements.append(Spacer(1, 20))
                
                # Course Objective
                elements.append(Paragraph('Course Objective', styles['Heading1']))
                elements.append(Paragraph(course.course_objective, styles['Normal']))
                elements.append(Spacer(1, 12))
                
                # Course Outcome
                elements.append(Paragraph('Course Outcome', styles['Heading1']))
                elements.append(Paragraph(course.course_outcome, styles['Normal']))
                elements.append(Spacer(1, 12))
                
                # Enrollment Details
                elements.append(Paragraph('Enrollment Details', styles['Heading1']))
                elements.append(Paragraph(f"Total Students Enrolled: {enrollments.count()}", styles['Normal']))
                elements.append(Spacer(1, 12))
                
                # Student List
                elements.append(Paragraph('List of Students', styles['Heading1']))
                
                # Student table data
                student_data = [['Student Name', 'Registration Number', 'Attendance', 'Certificate']]
                
                for enrollment in enrollments:
                    # Calculate attendance
                    total_days = AttendanceRecord.objects.filter(enrollment=enrollment).count()
                    present_days = AttendanceRecord.objects.filter(enrollment=enrollment, is_present=True).count()
                    attendance_percentage = (present_days / total_days) * 100 if total_days > 0 else 0
                    
                    student_data.append([
                        enrollment.student.name,
                        enrollment.student.registration_number,
                        f"{present_days}/{total_days} ({attendance_percentage:.1f}%)",
                        'Yes' if hasattr(enrollment, 'certificate') else 'No'
                    ])
                
                student_table = Table(student_data, colWidths=[120, 120, 120, 100])
                student_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                elements.append(student_table)
                elements.append(Spacer(1, 20))
                
                # Generation date
                elements.append(Paragraph(f"Report Generated On: {timezone.now().date()}", styles['Normal']))
                elements.append(Spacer(1, 30))
                
                # Signature fields
                signature_data = [
                    ['_'*20, '_'*20, '_'*20],
                    ['Course Coordinator', 'HOD', 'Principal']
                ]
                
                signature_table = Table(signature_data, colWidths=[150, 150, 150])
                signature_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                
                elements.append(signature_table)
                
                # Build the PDF
                doc.build(elements)
                
                # Save the report to CourseReport model
                try:
                    report = course.report
                except CourseReport.DoesNotExist:
                    report = CourseReport(course=course)
                
                with open(temp_path, 'rb') as f:
                    report_filename = f"report_{course.course_code}.pdf"
                    report.report_file.save(report_filename, BytesIO(f.read()))
                
                # Clean up the temporary file
                os.unlink(temp_path)
                
                messages.success(request, 'Report generated successfully as PDF.')
                
                # Return the file as a download
                response = HttpResponse(content_type='application/pdf')
                response['Content-Disposition'] = f'attachment; filename="{report_filename}"'
                with open(report.report_file.path, 'rb') as f:
                    response.write(f.read())
                return response
                
            except ImportError:
                # Fallback to plain text if ReportLab is not installed
                messages.warning(request, 'ReportLab library not installed. Generating text report instead.')
                
                # Create report content as text (your original text generation code)
                report_content = f"""
                CERTIFICATE COURSE REPORT
                
                Course Name: {course.name}
                Course Code: {course.course_code}
                Department: {course.department.name}
                Duration: {course.start_date} to {course.end_date}
                Total Hours: {course.total_hours}
                
                COURSE OBJECTIVE:
                {course.course_objective}
                
                COURSE OUTCOME:
                {course.course_outcome}
                
                ENROLLMENT DETAILS:
                Total Students Enrolled: {enrollments.count()}
                
                LIST OF STUDENTS:
                """
                
                for idx, enrollment in enumerate(enrollments, 1):
                    # Calculate attendance
                    total_days = AttendanceRecord.objects.filter(enrollment=enrollment).count()
                    present_days = AttendanceRecord.objects.filter(enrollment=enrollment, is_present=True).count()
                    attendance_percentage = (present_days / total_days) * 100 if total_days > 0 else 0
                    
                    report_content += f"""
                {idx}. {enrollment.student.name} (Reg No: {enrollment.student.registration_number})
                   Attendance: {present_days}/{total_days} ({attendance_percentage:.1f}%)
                   Certificate Issued: {'Yes' if hasattr(enrollment, 'certificate') else 'No'}
                    """
                
                report_content += f"""
                Report Generated On: {timezone.now().date()}
                """
                
                # Create a temporary file
                with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as temp_file:
                    temp_file.write(report_content.encode('utf-8'))
                    temp_path = temp_file.name
                
                # Save the report file to the model
                try:
                    report = course.report
                except CourseReport.DoesNotExist:
                    report = CourseReport(course=course)
                
                with open(temp_path, 'rb') as f:
                    report_filename = f"report_{course.course_code}.txt"
                    report.report_file.save(report_filename, BytesIO(f.read()))
                
                # Clean up the temporary file
                os.unlink(temp_path)
                
                messages.success(request, 'Report generated successfully as text file.')
                return redirect('view_report', report_pk=report.pk)
    
    context = {
        'course': course,
        'can_generate_report': can_generate_report,
        'enrollments': enrollments,
    }
    
    return render(request, 'certificate_courses/generate_report.html', context)


@login_required
def view_report(request, report_pk):
    department = request.user.department
    report = get_object_or_404(CourseReport, pk=report_pk, course__department=department)
    
    # In a real application, you would serve the report file
    # Here we'll just redirect to the report file URL
    return redirect(report.report_file.url)


# certificate_courses/views.py (add this)
import csv
from io import TextIOWrapper

@login_required
def import_students(request):
    department = request.user.department
    
    if request.method == 'POST':
        form = StudentBulkImportForm(request.POST, request.FILES)
        if form.is_valid():
            csv_file = request.FILES['csv_file']
            csv_file = TextIOWrapper(csv_file.file, encoding='utf-8')
            
            reader = csv.reader(csv_file)
            next(reader)  # Skip header row
            
            success_count = 0
            error_count = 0
            errors = []
            
            for row in reader:
                if len(row) >= 3:
                    name = row[0].strip()
                    reg_number = row[1].strip()
                    email = row[2].strip()
                    
                    # Check if student already exists
                    existing_student = Student.objects.filter(
                        registration_number=reg_number,
                        department=department
                    ).first()
                    
                    if existing_student:
                        errors.append(f"Student with registration number {reg_number} already exists.")
                        error_count += 1
                        continue
                    
                    # Create new student
                    try:
                        Student.objects.create(
                            name=name,
                            registration_number=reg_number,
                            email=email,
                            department=department
                        )
                        success_count += 1
                    except Exception as e:
                        errors.append(f"Error creating student {name} ({reg_number}): {str(e)}")
                        error_count += 1
                else:
                    error_count += 1
                    errors.append(f"Row does not have enough columns: {row}")
            
            if success_count > 0:
                messages.success(request, f'Successfully imported {success_count} students.')
            
            if error_count > 0:
                messages.warning(request, f'Failed to import {error_count} students.')
                for error in errors[:5]:  # Show first 5 errors
                    messages.error(request, error)
                
                if len(errors) > 5:
                    messages.warning(request, f"... and {len(errors) - 5} more errors.")
            
            return redirect('student_list')
    else:
        form = StudentBulkImportForm()
    
    return render(request, 'certificate_courses/import_students.html', {'form': form})



# certificate_courses/views.py (add this)
@login_required
def update_course_status(request, pk):
    department = request.user.department
    course = get_object_or_404(CertificateCourse, pk=pk, department=department)
    
    if request.method == 'POST':
        new_status = request.POST.get('status')
        if new_status in dict(CertificateCourse.STATUS_CHOICES).keys():
            course.status = new_status
            course.save()
            messages.success(request, f'Course status updated to {dict(CertificateCourse.STATUS_CHOICES)[new_status]}.')
        else:
            messages.error(request, 'Invalid status value.')
        
        return redirect('course_detail', pk=course.pk)
    
    return render(request, 'certificate_courses/update_status.html', {'course': course})

# certificate_courses/views.py (add this)
import csv

@login_required
def export_course_data(request, pk):
    department = request.user.department
    course = get_object_or_404(CertificateCourse, pk=pk, department=department)
    
    # Create the HttpResponse object with CSV header
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="{course.course_code}_data.csv"'
    
    writer = csv.writer(response)
    writer.writerow(['Student Name', 'Registration Number', 'Email', 'Enrollment Date', 'Attendance'])
    
    enrollments = course.enrollments.all()
    
    for enrollment in enrollments:
        student = enrollment.student
        
        # Calculate attendance
        total_days = AttendanceRecord.objects.filter(enrollment=enrollment).count()
        present_days = AttendanceRecord.objects.filter(enrollment=enrollment, is_present=True).count()
        attendance_percentage = (present_days / total_days) * 100 if total_days > 0 else 0
        
        writer.writerow([
            student.name,
            student.registration_number,
            student.email,
            enrollment.enrollment_date,
            f"{present_days}/{total_days} ({attendance_percentage:.1f}%)"
        ])
    
    return response

# certificate_courses/views.py (add these)
@login_required
def template_list(request):
    department = request.user.department
    templates = CertificateTemplate.objects.filter(department=department)
    return render(request, 'certificate_courses/template_list.html', {'templates': templates})

@login_required
def template_create(request):
    department = request.user.department
    
    if request.method == 'POST':
        form = CertificateTemplateForm(request.POST, request.FILES)
        if form.is_valid():
            template = form.save(commit=False)
            template.department = department
            template.save()
            messages.success(request, 'Certificate template created successfully.')
            return redirect('template_list')
    else:
        form = CertificateTemplateForm()
    
    return render(request, 'certificate_courses/template_form.html', {'form': form})

@login_required
def template_update(request, pk):
    department = request.user.department
    template = get_object_or_404(CertificateTemplate, pk=pk, department=department)
    
    if request.method == 'POST':
        form = CertificateTemplateForm(request.POST, request.FILES, instance=template)
        if form.is_valid():
            form.save()
            messages.success(request, 'Certificate template updated successfully.')
            return redirect('template_list')
    else:
        form = CertificateTemplateForm(instance=template)
    
    return render(request, 'certificate_courses/template_form.html', {'form': form, 'template': template})

@login_required
def template_delete(request, pk):
    department = request.user.department
    template = get_object_or_404(CertificateTemplate, pk=pk, department=department)
    
    if request.method == 'POST':
        template.delete()
        messages.success(request, 'Certificate template deleted successfully.')
        return redirect('template_list')
    
    return render(request, 'certificate_courses/template_confirm_delete.html', {'template': template})